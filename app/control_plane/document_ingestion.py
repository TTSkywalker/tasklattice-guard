from __future__ import annotations

import hashlib
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from docx import Document


MAX_DOCUMENTS = 3
MAX_DOCUMENT_BYTES = 5 * 1024 * 1024
MAX_TOTAL_BYTES = 10 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 120_000
MAX_DOCX_MEMBERS = 500
MAX_DOCX_UNCOMPRESSED_BYTES = 25 * 1024 * 1024
SUPPORTED_EXTENSIONS = (".doc", ".docx", ".txt")


class DocumentIngestionError(ValueError):
    pass


@dataclass(frozen=True, slots=True)
class DocumentSection:
    reference: str
    heading: str
    text: str


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    id: str
    name: str
    format: str
    size_bytes: int
    sha256: str
    character_count: int
    sections: tuple[DocumentSection, ...]

    def analysis_text(self) -> str:
        lines = [f"DOCUMENT {self.id}: {self.name}"]
        for section in self.sections:
            heading = f" | {section.heading}" if section.heading else ""
            lines.append(f"[SOURCE {section.reference}{heading}]")
            lines.append(section.text)
        return "\n".join(lines)

    def public_payload(self) -> dict[str, object]:
        return {
            "id": self.id,
            "name": self.name,
            "format": self.format,
            "size_bytes": self.size_bytes,
            "sha256": self.sha256,
            "character_count": self.character_count,
            "section_count": len(self.sections),
        }


def extract_documents(files: tuple[tuple[str, bytes], ...]) -> tuple[ExtractedDocument, ...]:
    if not files:
        raise DocumentIngestionError("Select at least one document to analyze.")
    if len(files) > MAX_DOCUMENTS:
        raise DocumentIngestionError(f"Upload no more than {MAX_DOCUMENTS} documents.")
    total_bytes = sum(len(content) for _, content in files)
    if total_bytes > MAX_TOTAL_BYTES:
        raise DocumentIngestionError("The combined document size exceeds 10 MB.")

    documents = tuple(
        extract_document(index=index, filename=filename, content=content)
        for index, (filename, content) in enumerate(files, start=1)
    )
    extracted_characters = sum(item.character_count for item in documents)
    if extracted_characters > MAX_EXTRACTED_CHARACTERS:
        raise DocumentIngestionError(
            "The documents contain too much text for one analysis. Split them into smaller files."
        )
    return documents


def extract_document(*, index: int, filename: str, content: bytes) -> ExtractedDocument:
    safe_name = Path(filename or "document").name.strip()
    extension = Path(safe_name).suffix.casefold()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentIngestionError(
            "Unsupported document type. Upload Word (.doc or .docx) or plain text (.txt)."
        )
    if not content:
        raise DocumentIngestionError(f"{safe_name} is empty.")
    if len(content) > MAX_DOCUMENT_BYTES:
        raise DocumentIngestionError(f"{safe_name} exceeds the 5 MB file limit.")

    document_id = f"document-{index}"
    if extension == ".docx":
        sections = _docx_sections(document_id, safe_name, content)
        format_name = "docx"
    elif extension == ".doc":
        sections = _legacy_doc_sections(document_id, safe_name, content)
        format_name = "doc"
    else:
        sections = _text_sections(document_id, safe_name, content)
        format_name = "txt"

    character_count = sum(len(item.text) for item in sections)
    if character_count < 20:
        raise DocumentIngestionError(
            f"{safe_name} does not contain enough readable text to analyze."
        )
    if character_count > MAX_EXTRACTED_CHARACTERS:
        raise DocumentIngestionError(
            f"{safe_name} contains too much text for one analysis. Split it into smaller files."
        )
    return ExtractedDocument(
        id=document_id,
        name=safe_name,
        format=format_name,
        size_bytes=len(content),
        sha256=hashlib.sha256(content).hexdigest(),
        character_count=character_count,
        sections=sections,
    )


def _docx_sections(
    document_id: str, filename: str, content: bytes
) -> tuple[DocumentSection, ...]:
    if not content.startswith(b"PK"):
        raise DocumentIngestionError(f"{filename} is not a valid DOCX document.")
    try:
        with zipfile.ZipFile(BytesIO(content)) as archive:
            members = archive.infolist()
            if len(members) > MAX_DOCX_MEMBERS:
                raise DocumentIngestionError(f"{filename} contains too many embedded files.")
            if any(item.flag_bits & 0x1 for item in members):
                raise DocumentIngestionError(f"Encrypted DOCX files are not supported: {filename}.")
            if sum(item.file_size for item in members) > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise DocumentIngestionError(f"{filename} expands beyond the safe DOCX limit.")
        document = Document(BytesIO(content))
    except DocumentIngestionError:
        raise
    except (OSError, ValueError, zipfile.BadZipFile) as error:
        raise DocumentIngestionError(f"{filename} could not be read as DOCX.") from error

    sections: list[DocumentSection] = []
    current_heading = ""
    paragraph_index = 0
    for paragraph in document.paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        paragraph_index += 1
        style_name = (paragraph.style.name if paragraph.style else "").casefold()
        if style_name.startswith("heading") or style_name in {"title", "subtitle"}:
            current_heading = text
        sections.append(
            DocumentSection(
                reference=f"{document_id}:paragraph-{paragraph_index}",
                heading=current_heading,
                text=text,
            )
        )
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            values = [_clean_text(cell.text) for cell in row.cells]
            text = " | ".join(value for value in values if value)
            if text:
                sections.append(
                    DocumentSection(
                        reference=f"{document_id}:table-{table_index}-row-{row_index}",
                        heading=current_heading,
                        text=text,
                    )
                )
    return tuple(sections)


def _legacy_doc_sections(
    document_id: str, filename: str, content: bytes
) -> tuple[DocumentSection, ...]:
    if not content.startswith(bytes.fromhex("D0CF11E0A1B11AE1")):
        raise DocumentIngestionError(f"{filename} is not a valid legacy Word document.")
    executable = shutil.which("antiword")
    converter = "antiword"
    if executable is None and os.uname().sysname == "Darwin":
        executable = shutil.which("textutil")
        converter = "textutil"
    if executable is None:
        raise DocumentIngestionError(
            "Legacy .doc conversion is unavailable on this server. Save the file as .docx and retry."
        )
    try:
        with tempfile.TemporaryDirectory(prefix="tasklattice-doc-") as directory:
            source = Path(directory) / "source.doc"
            source.write_bytes(content)
            command = (
                [executable, str(source)]
                if converter == "antiword"
                else [executable, "-convert", "txt", "-stdout", str(source)]
            )
            result = subprocess.run(
                command,
                cwd=directory,
                stdin=subprocess.DEVNULL,
                capture_output=True,
                check=False,
                timeout=10,
            )
    except (OSError, subprocess.SubprocessError) as error:
        raise DocumentIngestionError(f"{filename} could not be converted safely.") from error
    if result.returncode != 0 or not result.stdout.strip():
        raise DocumentIngestionError(f"{filename} could not be read as a Word document.")
    return _text_sections(document_id, filename, result.stdout)


def _text_sections(
    document_id: str, filename: str, content: bytes
) -> tuple[DocumentSection, ...]:
    text = _decode_text(filename, content)
    lines = text.splitlines()
    sections: list[DocumentSection] = []
    for start in range(0, len(lines), 20):
        chunk = _clean_text("\n".join(lines[start : start + 20]))
        if not chunk:
            continue
        end = min(start + 20, len(lines))
        sections.append(
            DocumentSection(
                reference=f"{document_id}:lines-{start + 1}-{end}",
                heading="",
                text=chunk,
            )
        )
    return tuple(sections)


def _decode_text(filename: str, content: bytes) -> str:
    encodings = ("utf-8-sig", "utf-16", "gb18030")
    for encoding in encodings:
        try:
            value = content.decode(encoding)
        except UnicodeDecodeError:
            continue
        if "\x00" not in value:
            return value
    raise DocumentIngestionError(f"{filename} is not valid UTF-8, UTF-16, or GB18030 text.")


def _clean_text(value: str) -> str:
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in value.splitlines()]
    return "\n".join(line for line in lines if line).strip()
