from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document

from app.control_plane.document_ingestion import (
    DocumentIngestionError,
    extract_document,
    extract_documents,
)


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Data handling", level=1)
    document.add_paragraph(
        "Customer-service staff may use the assistant for approved account support."
    )
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Restricted"
    table.cell(0, 1).text = "Never disclose account credentials"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_docx_extraction_preserves_headings_tables_and_source_references():
    item = extract_document(
        index=1,
        filename="customer-policy.docx",
        content=_docx_bytes(),
    )

    assert item.format == "docx"
    assert item.sections[0].reference == "document-1:paragraph-1"
    assert item.sections[1].heading == "Data handling"
    assert any("Never disclose" in section.text for section in item.sections)
    assert "[SOURCE document-1:table-1-row-1" in item.analysis_text()


def test_plain_text_extraction_supports_chinese_and_bounded_line_references():
    item = extract_document(
        index=2,
        filename="合规要求.txt",
        content="第一条：只允许客服分析。\n第二条：禁止披露个人信息。".encode("utf-8"),
    )

    assert item.format == "txt"
    assert item.sections[0].reference == "document-2:lines-1-2"
    assert "禁止披露个人信息" in item.sections[0].text


def test_legacy_doc_uses_bounded_converter(monkeypatch):
    monkeypatch.setattr(
        "app.control_plane.document_ingestion.shutil.which",
        lambda name: "/usr/bin/antiword" if name == "antiword" else None,
    )
    monkeypatch.setattr(
        "app.control_plane.document_ingestion.subprocess.run",
        lambda *args, **kwargs: SimpleNamespace(
            returncode=0,
            stdout=b"Legacy compliance policy permits support and blocks credential disclosure.",
        ),
    )
    item = extract_document(
        index=1,
        filename="legacy.doc",
        content=bytes.fromhex("D0CF11E0A1B11AE1") + b"legacy-content",
    )

    assert item.format == "doc"
    assert "blocks credential disclosure" in item.sections[0].text


def test_document_ingestion_rejects_pdf_empty_and_more_than_three_files():
    with pytest.raises(DocumentIngestionError, match="Unsupported"):
        extract_document(index=1, filename="policy.pdf", content=b"%PDF-1.7")
    with pytest.raises(DocumentIngestionError, match="empty"):
        extract_document(index=1, filename="policy.txt", content=b"")
    with pytest.raises(DocumentIngestionError, match="no more than 3"):
        extract_documents(
            tuple(
                (f"policy-{index}.txt", b"A sufficiently long policy document.")
                for index in range(4)
            )
        )
    with pytest.raises(DocumentIngestionError, match="too much text"):
        extract_document(
            index=1,
            filename="oversized-text.txt",
            content=("reviewed compliance requirement\n" * 5_000).encode("utf-8"),
        )
